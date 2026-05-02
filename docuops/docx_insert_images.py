"""
docx_insert_images.py — Insert Images into a DOCX Document
============================================================
Provides a helper for adding a heading, image, and page break into an
existing ``python-docx`` Document object.  The caller is responsible for
creating the Document and calling ``doc.save()``.

Public API
----------
    insert_image(doc, img_path, size_x, size_y, title=None)

Dependencies
------------
    python-docx (pip, imported as ``docx``)
"""

import docx as _docx
from docx.shared import Inches


def insert_image(
    doc: _docx.Document,
    img_path: str,
    size_x: float,
    size_y: float,
    title: str | None = None,
) -> None:
    """Insert an image (with an optional heading) into *doc*, then add a page break.

    Args:
        doc: An open ``python-docx`` Document object to write into.
        img_path: Path to the image file to embed.
        size_x: Image width in inches.
        size_y: Image height in inches.
        title: Optional heading text to add before the image.  If ``None``,
            no heading is inserted.

    Example::

        import docx
        from docuops.docx_insert_images import insert_image

        doc = docx.Document()
        insert_image(doc, "logo.jpg", 5, 7, title="Business License")
        insert_image(doc, "gst.jpg",  5, 7, title="GST")
        doc.save("output.docx")
    """
    if title is not None:
        doc.add_heading(title, level=0)
    doc.add_picture(img_path, width=Inches(size_x), height=Inches(size_y))
    doc.add_page_break()
